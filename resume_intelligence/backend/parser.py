"""Resume parser for PDF and DOCX files.

Text extraction:
- PDF: PyMuPDF primary (https://pymupdf.readthedocs.io/)
- PDF fallback: pdfplumber (https://github.com/jsvine/pdfplumber)
- DOCX: python-docx
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass
from typing import Any, Optional

import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from models.resume import ParseResult, ResumeData
from utils.file_utils import validate_resume_file
from utils.logger import get_logger

logger = get_logger(__name__)

EMAIL_PATTERN = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b")
URL_PATTERN = re.compile(
    r"https?://(?:www\.)?(?:linkedin\.com|github\.com|[\w.-]+)/[\w./?=#&%-]+",
    re.IGNORECASE,
)
BULLET_PATTERN = re.compile(r"^[\u2022\u2023\u25E6\u2043\u2219\-\*\•]\s*")

SECTION_HEADERS = {
    "education": [
        "education",
        "academic background",
        "academics",
        "qualification",
        "qualifications",
    ],
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "technologies",
        "tech stack",
    ],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "work history",
    ],
    "projects": [
        "projects",
        "personal projects",
        "academic projects",
        "key projects",
    ],
    "links": [
        "links",
        "profiles",
        "contact",
    ],
}

# Lines that look like section headers, not names
HEADER_SKIP_WORDS = set()
for headers in SECTION_HEADERS.values():
    HEADER_SKIP_WORDS.update(headers)


@dataclass
class SectionBlock:
    """A detected resume section with its raw lines."""

    name: str
    lines: list[str]


def parse_resume(file_bytes: bytes, filename: str) -> ParseResult:
    """
    Parse a resume file into structured ResumeData.

    Returns partial data with warnings when parsing is incomplete.
    """
    validation = validate_resume_file(filename, file_bytes)
    if not validation.valid:
        return ParseResult(success=False, errors=[validation.error or "Invalid file."])

    warnings: list[str] = []
    errors: list[str] = []

    try:
        raw_text = _extract_text(file_bytes, validation.extension)
    except Exception as exc:
        logger.error("Text extraction failed for %s: %s", filename, exc)
        return ParseResult(success=False, errors=[f"Failed to extract text: {exc}"])

    if not raw_text.strip():
        return ParseResult(
            success=False,
            errors=["No readable text found in the document."],
            data=ResumeData(raw_text="", parse_warnings=["Empty document."]),
        )

    lines = _normalize_lines(raw_text)
    sections = _detect_sections(lines)

    email = _extract_email(raw_text)
    links = _extract_links(raw_text)
    name = _extract_name(lines, email)

    skills = _parse_skills(sections.get("skills", []))
    education = _parse_entries(sections.get("education", []), entry_type="education")
    experience = _parse_entries(sections.get("experience", []), entry_type="experience")
    projects = _parse_entries(sections.get("projects", []), entry_type="projects")

    if not email:
        warnings.append("Email not found in resume.")
    if not skills:
        warnings.append("Skills section not detected or empty.")
    if not experience and not projects:
        warnings.append("No experience or projects section detected.")

    for section_name in ("education", "skills", "experience", "projects"):
        if section_name not in sections:
            warnings.append(f"Section '{section_name}' not detected.")

    data = ResumeData(
        name=name,
        email=email,
        education=education,
        skills=skills,
        projects=projects,
        experience=experience,
        links=links,
        raw_text=raw_text,
        parse_warnings=warnings,
    )

    success = bool(email or skills or experience or projects or education)
    if not success:
        errors.append("Could not extract meaningful resume content.")

    return ParseResult(success=success, data=data, errors=errors)


def _extract_text(file_bytes: bytes, extension: str) -> str:
    """Extract plain text from PDF or DOCX."""
    if extension == ".pdf":
        return _extract_pdf_text(file_bytes)
    if extension == ".docx":
        return _extract_docx_text(file_bytes)
    raise ValueError(f"Unsupported extension: {extension}")


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Try PyMuPDF first; fall back to pdfplumber for sparse layouts."""
    text = _extract_with_pymupdf(file_bytes)

    # Sparse extraction often means multi-column or scanned layout issues
    if len(text.strip()) < 100:
        logger.info("PyMuPDF yielded sparse text; trying pdfplumber fallback.")
        fallback = _extract_with_pdfplumber(file_bytes)
        if len(fallback.strip()) > len(text.strip()):
            return fallback

    return text


def _extract_with_pymupdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        pages = [page.get_text("text") for page in doc]
        return "\n".join(pages)
    finally:
        doc.close()


def _extract_with_pdfplumber(file_bytes: bytes) -> str:
    import io

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages)


def _extract_docx_text(file_bytes: bytes) -> str:
    import io

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _normalize_lines(text: str) -> list[str]:
    """Split text into clean non-empty lines."""
    lines = []
    for line in text.splitlines():
        cleaned = line.strip()
        if cleaned:
            lines.append(cleaned)
    return lines


def _is_section_header(line: str) -> Optional[str]:
    """Return section key if line looks like a section header."""
    normalized = re.sub(r"[^a-z\s]", "", line.lower()).strip()
    if not normalized or len(normalized) > 60:
        return None

    for section_key, headers in SECTION_HEADERS.items():
        for header in headers:
            if normalized == header or normalized.startswith(header + " "):
                return section_key
    return None


def _detect_sections(lines: list[str]) -> dict[str, list[str]]:
    """Group lines under detected section headers."""
    sections: dict[str, list[str]] = {}
    current_section: Optional[str] = None

    for line in lines:
        header = _is_section_header(line)
        if header:
            current_section = header
            sections.setdefault(current_section, [])
            continue

        if current_section:
            sections[current_section].append(line)

    return sections


def _extract_email(text: str) -> str:
    match = EMAIL_PATTERN.search(text)
    return match.group(0) if match else ""


def _extract_links(text: str) -> list[str]:
    found = URL_PATTERN.findall(text)
    # Deduplicate while preserving order
    seen: set[str] = set()
    links: list[str] = []
    for link in found:
        if link not in seen:
            seen.add(link)
            links.append(link)
    return links


def _extract_name(lines: list[str], email: str) -> str:
    """Heuristic: first substantial line that is not contact info or a header."""
    for line in lines[:8]:
        lower = line.lower()
        if email and email in line:
            continue
        if URL_PATTERN.search(line):
            continue
        if re.search(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", line):
            continue
        if _is_section_header(line):
            continue
        if lower in HEADER_SKIP_WORDS:
            continue
        # Skip lines that are mostly symbols or too long for a name
        if len(line) > 60 or len(line) < 2:
            continue
        if sum(c.isalpha() for c in line) < len(line) * 0.5:
            continue
        return line.strip()
    return ""


def _parse_skills(lines: list[str]) -> list[str]:
    """Parse skills from bullet or comma-separated lists."""
    if not lines:
        return []

    text = " ".join(lines)
    # Flatten inline lists like "Python, Java, SQL"
    parts = re.split(r"[,;|/]", text)
    skills: list[str] = []

    for part in parts:
        cleaned = BULLET_PATTERN.sub("", part).strip()
        if cleaned and len(cleaned) < 50:
            skills.append(cleaned)

    # Also pick up line-by-line bullets
    for line in lines:
        cleaned = BULLET_PATTERN.sub("", line).strip()
        if cleaned and cleaned not in skills and len(cleaned) < 50:
            if "," not in cleaned:
                skills.append(cleaned)

    return skills


def _parse_entries(lines: list[str], entry_type: str) -> list[dict[str, Any]]:
    """
    Parse education/experience/project blocks from section lines.

    Each entry is a dict with title, org, dates, description keys.
    """
    if not lines:
        return []

    entries: list[dict[str, Any]] = []
    current: dict[str, Any] = {}
    description_lines: list[str] = []

    date_pattern = re.compile(
        r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|"
        r"January|February|March|April|June|July|August|September|October|November|December)"
        r"[\w\s,\-–/]*\d{4}\b|\b\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)\b",
        re.IGNORECASE,
    )

    def flush_entry() -> None:
        nonlocal current, description_lines
        if current or description_lines:
            current["description"] = " ".join(description_lines).strip()
            if current.get("title") or current.get("description"):
                entries.append(current)
        current = {}
        description_lines = []

    for line in lines:
        is_bullet = bool(BULLET_PATTERN.match(line))
        cleaned = BULLET_PATTERN.sub("", line).strip()
        has_date = bool(date_pattern.search(cleaned))

        # New entry heuristic: non-bullet title line or line with dates
        if not is_bullet and (has_date or (len(cleaned) < 80 and cleaned.isupper() is False)):
            if current or description_lines:
                flush_entry()
            current = _split_title_org_dates(cleaned, entry_type)
            continue

        if is_bullet or description_lines or current:
            description_lines.append(cleaned)

    flush_entry()
    return entries


def _split_title_org_dates(line: str, entry_type: str) -> dict[str, Any]:
    """Split a header line into title, org, and dates where possible."""
    date_pattern = re.compile(
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|"
        r"January|February|March|April|June|July|August|September|October|November|December)"
        r"[\w\s,\-–/]*\d{4}|\d{4}\s*[-–]\s*(?:\d{4}|Present|Current))",
        re.IGNORECASE,
    )
    dates = date_pattern.findall(line)
    remainder = date_pattern.sub("", line).strip(" |-–,")

    parts = [p.strip() for p in re.split(r"\s+[-@|]\s+|\s+at\s+", remainder, maxsplit=1, flags=re.IGNORECASE)]

    entry: dict[str, Any] = {
        "title": parts[0] if parts else remainder,
        "org": parts[1] if len(parts) > 1 else "",
        "dates": dates[0] if dates else "",
        "description": "",
    }

    if entry_type == "education" and not entry["org"] and len(parts) == 1:
        # "B.Tech Computer Science, IIT Delhi" style
        school_split = re.split(r",\s*", remainder, maxsplit=1)
        if len(school_split) == 2:
            entry["title"] = school_split[0]
            entry["org"] = school_split[1]

    return entry


def parse_resume_from_path(path: str) -> ParseResult:
    """Parse a resume from a filesystem path (used in tests)."""
    with open(path, "rb") as handle:
        file_bytes = handle.read()
    filename = os.path.basename(path)
    return parse_resume(file_bytes, filename)
